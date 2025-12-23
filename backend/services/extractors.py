"""Document text extractors for various file formats."""

import os
from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import List, Iterator, Optional
from pathlib import Path


@dataclass
class TextBlock:
    """A block of text extracted from a document."""
    text: str
    page: Optional[int] = None
    block_index: int = 0
    metadata: Optional[dict] = None


class BaseExtractor(ABC):
    """Base class for document text extraction."""
    
    @abstractmethod
    def extract(self, file_path: str) -> Iterator[TextBlock]:
        """Extract text blocks from document."""
        pass
    
    @abstractmethod
    def get_total_pages(self, file_path: str) -> int:
        """Get total number of pages/sections."""
        pass


class PDFExtractor(BaseExtractor):
    """Extract text from PDF files using pdfplumber."""
    
    def extract(self, file_path: str) -> Iterator[TextBlock]:
        import pdfplumber
        
        block_index = 0
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                
                # Split into paragraphs
                paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
                
                for para in paragraphs:
                    if len(para) > 10:  # Skip very short blocks
                        yield TextBlock(
                            text=para,
                            page=page_num,
                            block_index=block_index,
                        )
                        block_index += 1
    
    def get_total_pages(self, file_path: str) -> int:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            return len(pdf.pages)


class DOCXExtractor(BaseExtractor):
    """Extract text from DOCX files using python-docx."""
    
    def extract(self, file_path: str) -> Iterator[TextBlock]:
        from docx import Document
        
        doc = Document(file_path)
        block_index = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) > 10:
                yield TextBlock(
                    text=text,
                    page=None,  # DOCX doesn't have page numbers
                    block_index=block_index,
                )
                block_index += 1
    
    def get_total_pages(self, file_path: str) -> int:
        from docx import Document
        doc = Document(file_path)
        return len(doc.paragraphs)


class TXTExtractor(BaseExtractor):
    """Extract text from plain text files."""
    
    def extract(self, file_path: str) -> Iterator[TextBlock]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        
        for i, para in enumerate(paragraphs):
            if len(para) > 10:
                yield TextBlock(
                    text=para,
                    page=None,
                    block_index=i,
                )
    
    def get_total_pages(self, file_path: str) -> int:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return len([p for p in content.split("\n\n") if p.strip()])


class MarkdownExtractor(TXTExtractor):
    """Extract text from Markdown files (same as TXT)."""
    pass


class CSVExtractor(BaseExtractor):
    """Extract text from CSV files."""
    
    def extract(self, file_path: str) -> Iterator[TextBlock]:
        import pandas as pd
        
        df = pd.read_csv(file_path, encoding="utf-8", on_bad_lines="skip")
        block_index = 0
        
        for row_idx, row in df.iterrows():
            for col_idx, value in enumerate(row):
                if pd.notna(value) and isinstance(value, str) and len(value) > 10:
                    yield TextBlock(
                        text=str(value),
                        page=None,
                        block_index=block_index,
                        metadata={"row": row_idx, "column": df.columns[col_idx]},
                    )
                    block_index += 1
    
    def get_total_pages(self, file_path: str) -> int:
        import pandas as pd
        df = pd.read_csv(file_path, encoding="utf-8", on_bad_lines="skip")
        return len(df)


def get_extractor(file_type: str) -> BaseExtractor:
    """Get appropriate extractor for file type."""
    extractors = {
        "pdf": PDFExtractor(),
        "docx": DOCXExtractor(),
        "doc": DOCXExtractor(),  # Requires conversion first
        "txt": TXTExtractor(),
        "md": MarkdownExtractor(),
        "csv": CSVExtractor(),
    }
    
    extractor = extractors.get(file_type.lower())
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")
    
    return extractor


def detect_file_type(filename: str) -> str:
    """Detect file type from extension."""
    ext = Path(filename).suffix.lower().lstrip(".")
    
    type_map = {
        "pdf": "pdf",
        "docx": "docx",
        "doc": "doc",
        "txt": "txt",
        "md": "md",
        "markdown": "md",
        "csv": "csv",
    }
    
    file_type = type_map.get(ext)
    if not file_type:
        raise ValueError(f"Unsupported file extension: {ext}")
    
    return file_type
